import os
import json
import base64
import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI
from pathlib import Path
from dotenv import load_dotenv

load_dotenv(Path(__file__).parent.resolve() / ".env")
client_openai = OpenAI()

def extract_text_from_file(file_path: str) -> str:
    ext = file_path.lower()
    text = ""
    if ext.endswith(".pdf"):
        try:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except: pass
    elif ext.endswith(".docx") or ext.endswith(".doc"):
        try:
            doc = Document(file_path)
            paragraphs_text = "\n".join([p.text for p in doc.paragraphs])
            tables_text = "\n".join([p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs])
            text = paragraphs_text + "\n" + tables_text
        except: pass
    return text

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract raw text directly from an in-memory byte stream, bypassing the disk."""
    ext = filename.lower()
    text = ""
    try:
        if ext.endswith(".pdf"):
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                text += page.get_text()
            doc.close()
        elif ext.endswith(".docx"):
            doc = Document(io.BytesIO(file_bytes))
            paragraphs_text = "\n".join([p.text for p in doc.paragraphs])
            tables_text = "\n".join([p.text for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs])
            text = paragraphs_text + "\n" + tables_text
        elif ext.endswith(".doc"):
            return "ERROR_UNSUPPORTED_DOC"
    except Exception as e:
        print(f"Error extracting memory file: {e}")
        return f"ERROR_CRASH: {str(e)}"
    return text

def parse_cv_to_corporate_json(raw_text: str):
    """
    Given the raw text of a candidate CV, ask OpenAI to extract all
    relevant historical fields strictly matching the corporate MF format.
    """
    prompt = f"""
INSTRUCCIÓN CRÍTICA Y OBLIGATORIA:
Eres un experto analista de recursos humanos. Dada la siguiente extracción en texto plano de un Curriculum Vitae, debes extraer y organizar rígidamente la información en un objeto JSON puro.

MUY IMPORTANTE SOBRE "stack_tecnologico":
Lee TODO el historial de proyectos del candidato. Extrae ABSOLUTAMENTE TODAS las tecnologías, herramientas, lenguajes, frameworks, bases de datos y sistemas que se mencionen en CUALQUIERA de sus trabajos o proyectos.
NO RESUMAS. Todo lo que encuentres DEBE clasificarse dentro del objeto "stack_tecnologico". Si no lo haces, el documento quedará inválido.

Tu salida debe ser ÚNICAMENTE un JSON válido (sin marcas de markdown, ni ```json, solo el JSON raw).

Estructura JSON exacta requerida:
{{
  "nombre_candidato": "Nombre completo de la persona",
  "titulo_profesional": "Titular profesional principal basado en su experiencia más senior. Ej: Analista Programador Java Senior",
  "perfil_resumen": "Párrafo de 2-3 líneas en tono corporativo describiendo quién es este candidato, cuántos años de experiencia tiene y en qué se especializa. Ej: Desarrollador Java Senior con más de 10 años de experiencia en entornos bancarios...",
  "perfil_experiencia_en": ["Área o capacidad clave 1", "Área o capacidad clave 2", "máximo 5 ítems que resuman sus capacidades principales"],
  "areas_especializacion": ["Área funcional 1", "Área funcional 2", "sus áreas de especialización técnica o funcional (ej. Desarrollo backend, Integración de sistemas, Gestión de proyectos)"],
  "stack_tecnologico": {{
    "Lenguajes": "Ej. Java, Python, C#",
    "Frameworks / Librerías": "Ej. Spring Boot, React, Angular",
    "Bases de datos": "Ej. Oracle, SQL Server, MySQL",
    "Cloud / DevOps": "Ej. AWS, Azure, Docker, Kubernetes",
    "Herramientas corporativas": "Ej. O365, Jira, Confluence, Active Directory",
    "ERP / CRM": "Ej. SAP, Salesforce",
    "Metodologías": "Ej. Scrum, Kanban, ITIL",
    "Otros": "Cualquier otra tecnología relevante no encuadrada arriba"
  }},
  "instruccion_antiolvido": "Comprueba que has incluido MongoDB, O365, SAP, Kafka, PowerCenter, PLCs, Azure si aparecen en el texto. NO RESUMAS tecnologías.",
  "sectores_economicos": ["Sector donde ha trabajado 1", "Sector 2"],
  "proyectos_realizados": [
    {{
      "empresa": "Nombre de la empresa o cliente",
      "fechas": "Rango de fechas (ej. MARZO 2023 – PRESENTE)",
      "proyecto": "Título del proyecto o rol desempeñado",
      "objetivo": "Párrafo breve explicando qué se hizo",
      "niveles_intervencion": ["Tarea específica 1", "Tarea específica 2"],
      "entorno_tecnologico": "Tecnologías del proyecto separadas por coma"
    }}
  ],
  "formacion_academica": [
    "Titulación en Especialidad — Institución (Año)"
  ],
  "idiomas": [
    "Inglés: Nivel Intermedio",
    "Español: Nativo"
  ]
}}

REGLAS ESTRICTAS:
- Si un dato no existe en el CV, usa array vacío [] u objeto vacío {{}}.
- NUNCA inventes certificaciones, cursos ni conocimientos que no estén EXPLÍCITAMENTE en el CV.
- Proyectos en orden cronológico INVERSO (más reciente primero).
- Tono corporativo y profesional.
- No incluyas datos personales (DNI, dirección, edad).

TEXTO DEL CURRICULUM:
--------------------
{raw_text[:12000]}
"""

    response = client_openai.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a helpful HR data extraction assistant. You only respond in strictly parsable JSON format matching the schema requested."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.1
    )
    
    response_text = response.choices[0].message.content.strip()
    if response_text.startswith("```json"):
        response_text = response_text[7:]
    if response_text.endswith("```"):
        response_text = response_text[:-3]
        
    try:
        return json.loads(response_text)
    except Exception as e:
        print(f"Error parsing JSON from LLM: {e}")
        return None

def add_header_logo(doc):
    logo_path = Path(__file__).parent / "logo.png"
    if logo_path.exists():
        header = doc.sections[0].header
        htable = header.add_table(1, 1, Inches(6))
        htable.autofit = True
        hcell = htable.rows[0].cells[0]
        hc_p = hcell.paragraphs[0]
        hc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hrun = hc_p.add_run()
        hrun.add_picture(str(logo_path), width=Inches(2.0))

def add_footer(doc):
    footer = doc.sections[0].footer
    
    # Use a table to have left logo and right text
    p_orig = footer.paragraphs[0]
    p_orig.text = ""
    
    table = footer.add_table(1, 2, Inches(6))
    table.autofit = True
    cells = table.rows[0].cells
    
    p_left = cells[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Busca todas las fotos que se llamen applus algo (ej. "applus1.png", "applus2.png")
    applus_logos = sorted(Path(__file__).parent.glob("applus*.png"))
    if applus_logos:
        run_img = p_left.add_run()
        for logo_path in applus_logos:
            run_img.add_picture(str(logo_path), width=Inches(0.6))
            run_img.add_text("   ") # Pequeño espaciador entre logos
    
    p_right = cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # We can fake the page numbering in Word manually if needed, but keeping it simple
    run = p_right.add_run("Reproducción prohibida")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

def generate_corporate_cv_docx(candidate_row, parsed_data) -> io.BytesIO:
    doc = Document()
    
    # 1. Estilos Globales
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0, 0, 0)
    
    add_header_logo(doc)
    add_footer(doc)
    
    cand_name = candidate_row.get('nombre', '') if candidate_row else parsed_data.get('nombre_candidato', '')
    cand_title = candidate_row.get('titulo_profesional', '') if candidate_row else parsed_data.get('titulo_profesional', 'Perfil Profesional')
    
    # Candidate Full Name top left
    p_name = doc.add_paragraph()
    run_name = p_name.add_run(cand_name.title() if cand_name else "")
    run_name.font.size = Pt(9)
    run_name.font.italic = True
    run_name.font.color.rgb = RGBColor(128, 128, 128)
    
    # Spacing
    doc.add_paragraph()
    
    # MAIN TITLE (Role)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(cand_title)
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    
    doc.add_paragraph() # Spacer

    # Helper function for underlined headers with full-width bottom border
    def add_section_header(title_text):
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.font.size = Pt(12)
        
        # Add full-width bottom border (Underline stretches to margins)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5 pt
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
        p.paragraph_format.space_after = Pt(8)

    # Helper function for generic square bullets
    def add_bullet_list(items):
        for item in items:
            p = doc.add_paragraph()
            p.add_run("▪\t" + item)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(0.25))
            p.paragraph_format.space_after = Pt(3)

    # 1. Perfil profesional
    perfil_resumen = parsed_data.get("perfil_resumen", "")
    perfil_exp_en = parsed_data.get("perfil_experiencia_en", [])
    if perfil_resumen or perfil_exp_en:
        add_section_header("Perfil profesional")
        if perfil_resumen:
            p_res = doc.add_paragraph()
            r_res = p_res.add_run(perfil_resumen)
            r_res.font.bold = True
            p_res.paragraph_format.space_after = Pt(6)
        if perfil_exp_en:
            p_amp = doc.add_paragraph()
            p_amp.add_run("Amplia experiencia en:")
            p_amp.paragraph_format.space_after = Pt(4)
            add_bullet_list(perfil_exp_en)
        doc.add_paragraph()

    # 2. Áreas de Especialización
    areas = parsed_data.get("areas_especializacion", [])
    if areas:
        add_section_header("Áreas de Especialización")
        add_bullet_list(areas)
        doc.add_paragraph()

    # 3. Stack Tecnológico Relevante
    raw_stack = parsed_data.get("stack_tecnologico", {})
    stack = {k: v for k, v in raw_stack.items()
             if v and isinstance(v, str) and str(v).lower() not in ["none", "null", "n/d", "n/a", ""]}
    if stack:
        add_section_header("Stack Tecnológico Relevante")
        for k, v in stack.items():
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(2.0))
            r_bold = p.add_run(k + ": ")
            r_bold.font.bold = True
            p.add_run(str(v))
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
        doc.add_paragraph()

    # 4. Sectores Económicos
    sectores = parsed_data.get("sectores_economicos", [])
    if sectores:
        add_section_header("Sectores Económicos")
        add_bullet_list(sectores)
        doc.add_paragraph()

    # 5. Proyectos Realizados
    proyectos = parsed_data.get("proyectos_realizados", [])
    if proyectos:
        add_section_header("Proyectos Realizados:")

        for proj in proyectos:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            cells = table.rows[0].cells

            p_comp = cells[0].paragraphs[0]
            p_comp.paragraph_format.space_before = Pt(6)
            r_comp = p_comp.add_run("- " + proj.get("empresa", "PROYECTO INDEPENDIENTE").upper())
            r_comp.font.bold = True

            p_date = cells[1].paragraphs[0]
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_date = p_date.add_run(proj.get("fechas", "").upper())
            r_date.font.bold = True

            p_proj = doc.add_paragraph()
            r_pref = p_proj.add_run("Proyecto: ")
            r_pref.font.bold = True
            p_proj.add_run("- " + proj.get("proyecto", ""))
            p_proj.paragraph_format.space_after = Pt(6)

            if proj.get("objetivo"):
                p_obj = doc.add_paragraph()
                r_obj = p_obj.add_run("Objetivo: ")
                r_obj.font.bold = True
                p_obj.add_run(proj.get("objetivo", ""))

            proj_interv = proj.get("niveles_intervencion", [])
            if proj_interv:
                p_niv = doc.add_paragraph()
                r_niv = p_niv.add_run("Niveles de Intervención:")
                r_niv.font.bold = True
                add_bullet_list(proj_interv)

            entorno = proj.get("entorno_tecnologico", "")
            if isinstance(entorno, list):
                entorno = ", ".join(entorno)
            if entorno and str(entorno).lower() not in ["none", "null", "n/d", "n/a", ""]:
                p_ent = doc.add_paragraph()
                r_ent = p_ent.add_run("Entorno tecnológico: ")
                r_ent.font.bold = True
                p_ent.add_run(str(entorno))
                p_ent.paragraph_format.space_after = Pt(12)
            else:
                if proj_interv or proj.get("objetivo"):
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 6. Formación
    formacion = parsed_data.get("formacion_academica", [])
    if formacion:
        add_section_header("Formación:")
        for f in formacion:
            p = doc.add_paragraph()
            parts = f.split(" en ", 1)
            if len(parts) == 2:
                r1 = p.add_run(parts[0] + " en ")
                r1.font.bold = True
                p.add_run(parts[1])
            else:
                p.add_run(f)
            p.paragraph_format.space_after = Pt(4)
        doc.add_paragraph()

    # 7. Idiomas
    idiomas = parsed_data.get("idiomas", [])
    if idiomas:
        add_section_header("Idiomas:")
        add_bullet_list(idiomas)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
